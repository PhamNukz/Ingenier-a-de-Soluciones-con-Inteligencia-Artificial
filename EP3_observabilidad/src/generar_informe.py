"""
GENERADOR DEL INFORME WORD — EP3 Observabilidad (IE7, IE8, IE9)
===============================================================

Construye Informe_EP3_ISY0101.docx leyendo los datos reales de logs/hallazgos.json
e incrustando los gráficos de docs/img/. Mantiene el informe sincronizado con las
métricas: si regeneras los datos (run_evaluacion.py) y vuelves a correr el análisis,
basta re-ejecutar este script para actualizar el informe.

Uso:
    cd EP3_observabilidad/src
    python analisis_hallazgos.py     # genera logs/hallazgos.json (si no existe)
    python generar_graficos.py       # genera docs/img/*.png (si no existen)
    python generar_informe.py
"""

import json
import os
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.append(os.path.dirname(__file__))
from dataset_queries import resumen as resumen_dataset

_BASE = os.path.dirname(os.path.dirname(__file__))
RUTA_HALLAZGOS = os.path.join(_BASE, "logs", "hallazgos.json")
DIR_IMG = os.path.join(_BASE, "docs", "img")
SALIDA = os.path.join(_BASE, "Informe_EP3_ISY0101.docx")

AZUL = RGBColor(0x1F, 0x4E, 0x79)
GRIS = RGBColor(0x55, 0x65, 0x73)


def cargar_hallazgos() -> dict:
    if not os.path.exists(RUTA_HALLAZGOS):
        raise SystemExit("Falta logs/hallazgos.json. Corre antes: python analisis_hallazgos.py")
    with open(RUTA_HALLAZGOS, encoding="utf-8") as f:
        return json.load(f)


# ── helpers de formato ────────────────────────────────────────────────────────
def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)


def h(doc, texto, size=13):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = AZUL
    p.space_after = Pt(2)
    return p


def parrafo(doc, texto, size=10, italic=False, color=None, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(texto)
    run.font.size = Pt(size)
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def imagen(doc, nombre, caption, ancho_cm=15.0):
    ruta = os.path.join(DIR_IMG, nombre)
    if not os.path.exists(ruta):
        parrafo(doc, f"[Figura faltante: {nombre} — corre generar_graficos.py]",
                italic=True, color=GRIS)
        return
    doc.add_picture(ruta, width=Cm(ancho_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = GRIS
    cap.paragraph_format.space_after = Pt(6)


def tabla_kv(doc, filas, encabezados=None):
    n_cols = len(filas[0])
    t = doc.add_table(rows=0, cols=n_cols)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if encabezados:
        celdas = t.add_row().cells
        for i, e in enumerate(encabezados):
            celdas[i].text = ""
            run = celdas[i].paragraphs[0].add_run(e)
            run.bold = True
            run.font.size = Pt(9)
    for fila in filas:
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = ""
            run = celdas[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    return t


# Logo de la portada: usa el primero que exista (permite reemplazar por el oficial
# guardándolo como logo_duoc.jpg o .png en docs/img/).
LOGO = next(
    (os.path.join(DIR_IMG, f"logo_duoc{ext}")
     for ext in (".jpg", ".jpeg", ".png")
     if os.path.exists(os.path.join(DIR_IMG, f"logo_duoc{ext}"))),
    os.path.join(DIR_IMG, "logo_duoc.png"),
)


def _espacio(doc, n=1, size=12):
    """Inserta n párrafos vacíos como separación vertical."""
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.add_run(" ").font.size = Pt(size)


def portada(doc, *, integrantes, docente, asignatura, fecha, titulo, subtitulo):
    """Construye la página de portada (estilo DuocUC) y un salto de página."""
    # Logo institucional (reemplazable por el oficial)
    if os.path.exists(LOGO):
        doc.add_picture(LOGO, width=Cm(11))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _espacio(doc, 3, 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INFORME")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = AZUL

    _espacio(doc, 1, 6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo.upper())
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitulo)
    r.italic = True
    r.font.size = Pt(11)

    _espacio(doc, 8, 12)  # empuja el bloque de autor hacia la parte baja

    for etiqueta, valor in [("Alumnos", integrantes),
                            ("Profesor", docente),
                            ("Asignatura", asignatura)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{etiqueta}: ")
        r.bold = True
        r.font.size = Pt(12)
        r2 = p.add_run(valor)
        r2.font.size = Pt(12)

    _espacio(doc, 3, 12)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(fecha)
    r.font.size = Pt(12)

    doc.add_page_break()


# ── construcción del documento ────────────────────────────────────────────────
def construir(d: dict):
    doc = Document()
    set_base_style(doc)

    res = d["resumen"]
    lat = d["latencia"]
    rec = d["recursos"]
    err = d["errores"]
    rel = d["relevancia_por_categoria"]
    ds = resumen_dataset()
    es_sim = res.get("fuente_datos") == "simulado"

    # ── Página de portada (estilo DuocUC) ────────────────────────────────────
    portada(
        doc,
        integrantes="Benjamín Aravena   ·   Francisco Gómez",
        docente="Roberto Eduardo Vega Araneda",
        asignatura="ISY0101 — Ingeniería de Soluciones con IA",
        fecha="01 de julio de 2026",
        titulo="Implementación de Observabilidad en un Agente de IA",
        subtitulo="Agente Portuario EPV — Asistente RAG de normativas del Puerto de Valparaíso",
    )

    # ── Metadatos (inicio de página 2) ───────────────────────────────────────
    parrafo(
        doc,
        "RA3 — Observabilidad, Seguridad y Ética en Agentes de IA · Evaluación Parcial N°3 (35%) · "
        "Repositorio: github.com/PhamNukz/Ingenier-a-de-Soluciones-con-Inteligencia-Artificial "
        "(módulo EP3_observabilidad)",
        size=8.5, italic=True, color=GRIS, after=6,
    )

    if es_sim:
        parrafo(
            doc,
            "Nota de transparencia: las cifras y figuras de este informe provienen de una "
            "corrida con datos de demostración (fuente_datos = \"simulado\", semilla fija), "
            "calibrados al comportamiento real del sistema. La instrumentación es funcional "
            "con datos reales ejecutando run_evaluacion.py con el token de GitHub Models; el "
            "pipeline de análisis y dashboard es idéntico en ambos casos.",
            size=8.5, italic=True, color=GRIS, after=6,
        )

    # ── 1. Introducción ──────────────────────────────────────────────────────
    h(doc, "1. Introducción y objetivo")
    parrafo(
        doc,
        "El sistema observado es el Agente Portuario EPV (EP2): un agente ReAct construido "
        "con LangChain que combina Recuperación Aumentada por Generación (RAG) sobre una base "
        "vectorial ChromaDB con el modelo gpt-4o-mini (servido vía GitHub Models) y cuatro "
        "herramientas: consulta de normativas internas, evaluación de cumplimiento, generación "
        "de reportes y búsqueda externa en Wikipedia. El objetivo de esta evaluación es "
        "instrumentar el agente con observabilidad para medir su precisión, latencia, "
        "consistencia, frecuencia de errores y uso de recursos; analizar sus trazas de "
        "ejecución para detectar cuellos de botella y patrones de falla; visualizar el "
        "comportamiento en un dashboard; e integrar controles de seguridad y uso responsable, "
        "proponiendo finalmente recomendaciones de optimización fundamentadas en datos.",
    )

    # ── 2. Arquitectura de instrumentación ───────────────────────────────────
    h(doc, "2. Arquitectura de instrumentación (IE1, IE2)")
    parrafo(
        doc,
        "Se diseñó una capa de observabilidad que envuelve al agente sin modificar su código. "
        "Un wrapper (medir_consulta) orquesta, por cada consulta: (a) controles de seguridad; "
        "(b) un retrieval instrumentado que replica la configuración de EP2 (mismos embeddings "
        "MiniLM multilingüe, misma base, k=4) capturando las distancias de ChromaDB; y (c) la "
        "ejecución del agente con un callback handler de LangChain que captura automáticamente "
        "el uso de tokens y la latencia de cada llamada al LLM dentro del bucle ReAct. Cada "
        "consulta produce un registro JSON estructurado (un objeto por línea, formato JSON "
        "Lines) que habilita la trazabilidad y alimenta el análisis y el dashboard.",
    )
    tabla_kv(
        doc,
        encabezados=["Métrica", "Qué se mide", "Fuente / técnica"],
        filas=[
            ["Precisión/Relevancia", "Score de similitud de los chunks (1/(1+distancia))", "Distancias de ChromaDB"],
            ["Latencia", "Total y desglose embedding → retrieval → LLM", "perf_counter + callback"],
            ["Consistencia", "Variación de respuesta/tokens al repetir la consulta", "N=5 repeticiones por query"],
            ["Frecuencia de errores", "Timeouts, errores de API, respuestas vacías, bloqueos", "Excepciones + callback on_llm_error"],
            ["Uso de recursos", "Tokens prompt+completion y nº de chunks", "token_usage del LLM"],
        ],
    )

    # ── 3. Metodología ───────────────────────────────────────────────────────
    h(doc, "3. Metodología y dataset (IE3)")
    cat = ds["por_categoria"]
    parrafo(
        doc,
        f"Se construyó un dataset de {ds['total_queries']} consultas variadas con "
        f"{ds['casos_limite']} casos límite, distribuidas en cuatro categorías: en-dominio "
        f"({cat.get('en_dominio',0)}), ambiguas ({cat.get('ambigua',0)}), fuera-dominio "
        f"({cat.get('fuera_dominio',0)}) y adversariales ({cat.get('adversarial',0)}). "
        f"Las consultas marcadas para consistencia se ejecutaron {ds['n_consistencia']} veces. "
        f"En total se registraron {res['total_ejecuciones']} ejecuciones "
        f"({res['exitosas']} exitosas, {res['con_error']} con error). Las preguntas ambiguas y "
        f"fuera-de-dominio estresan la cobertura del corpus; las adversariales prueban la capa "
        f"de seguridad frente a inyección de prompts.",
    )

    # ── 4. Resultados y análisis ─────────────────────────────────────────────
    h(doc, "4. Resultados y análisis (IE1–IE4)")
    pf = lat["porcentaje_tiempo_por_fase"]
    dpr = lat["desglose_promedio"]
    tabla_kv(
        doc,
        encabezados=["Indicador", "Valor", "Indicador", "Valor"],
        filas=[
            ["Latencia media", f"{lat['promedio_total_s']} s", "Latencia p95", f"{lat['p95_total_s']} s"],
            ["Latencia mediana", f"{lat['mediana_total_s']} s", "Latencia p99", f"{lat['p99_total_s']} s"],
            ["Tokens medios", f"{rec['tokens_promedio_total']}", "Tokens máx.", f"{rec['tokens_max_total']}"],
            ["Tasa error global", f"{err['tasa_error_global_pct']} %", "Tasa error API", f"{err['tasa_error_api_pct']} %"],
            ["Bloqueos seguridad", f"{err['bloqueados_seguridad']}", "Pasos promedio", f"{res['pasos_promedio']}"],
        ],
    )

    parrafo(doc, "4.1 Latencia y cuello de botella (IE2)", size=10.5).runs[0].bold = True
    parrafo(
        doc,
        f"El desglose por fase muestra que la generación con el LLM concentra el {pf['llm_pct']}% "
        f"del tiempo ({dpr['llm_s']} s en promedio), frente a un {pf['retrieval_pct']}% del "
        f"retrieval ({dpr['retrieval_s']} s) y un {pf['embedding_pct']}% del embedding "
        f"({dpr['embedding_s']} s). El cuello de botella es inequívocamente el LLM: optimizar el "
        f"RAG aportaría una mejora marginal. La brecha entre la mediana "
        f"({lat['mediana_total_s']} s) y el p95 ({lat['p95_total_s']} s, p99 {lat['p99_total_s']} s) "
        f"la explican las consultas con más pasos ReAct (más llamadas al LLM) y los fallos de API.",
    )
    imagen(doc, "02_desglose_latencia.png",
           "Figura 1. Desglose de latencia promedio por fase del pipeline RAG.", 15.0)
    imagen(doc, "01_latencia_temporal.png",
           "Figura 2. Latencia total por consulta con líneas de media y p95.", 15.0)

    parrafo(doc, "4.2 Uso de recursos / tokens (IE2)", size=10.5).runs[0].bold = True
    ratio = rec["ratio_prompt_completion"]
    if ratio >= 1.2:
        nota_ratio = (
            "El prompt domina el costo porque el formato ReAct reinyecta el scratchpad "
            "(razonamiento + observaciones) en cada paso, y crece con el número de pasos."
        )
    else:
        nota_ratio = (
            "La proporción prompt:completion es cercana a 1, propia de respuestas extensas "
            "sobre prompts compactos; el costo crece con el número de pasos ReAct."
        )
    nota_estim = ""
    if rec.get("consultas_tokens_estimados"):
        nota_estim = (
            f" Nota: GitHub Models no siempre devuelve token_usage, por lo que "
            f"{rec['consultas_tokens_estimados']} consultas se contabilizaron con un respaldo "
            f"local (tiktoken), garantizando cobertura completa de la métrica."
        )
    parrafo(
        doc,
        f"El consumo promedio es de {rec['tokens_promedio_total']} tokens por consulta "
        f"({rec['tokens_promedio_prompt']} de prompt y {rec['tokens_promedio_completion']} de "
        f"completion), con una razón prompt:completion de {ratio}:1. {nota_ratio}{nota_estim}",
    )
    imagen(doc, "03_distribucion_tokens.png",
           "Figura 3. Distribución de tokens por consulta y promedio prompt vs. completion.", 15.0)

    parrafo(doc, "4.3 Precisión/relevancia y errores (IE1, IE4)", size=10.5).runs[0].bold = True
    rel_ed = rel.get("en_dominio", {}).get("sim_top_promedio", "—")
    rel_fd = rel.get("fuera_dominio", {}).get("sim_top_promedio", "—")
    umbral = d.get("umbral_relevancia", "—")
    factor = round(rel_ed / rel_fd, 1) if isinstance(rel_ed, (int, float)) and rel_fd else "—"
    parrafo(
        doc,
        f"El score de similitud (derivado de la distancia L2 de ChromaDB, por lo que lo relevante "
        f"es la separación relativa y no el valor absoluto) distingue las categorías: en-dominio "
        f"promedia {rel_ed} frente a {rel_fd} en fuera-dominio (~{factor}× más relevante). Esta "
        f"brecha es un patrón accionable (IE4): un umbral de relevancia en torno a {umbral} "
        f"permitiría detectar consultas no cubiertas por el corpus y activar un fallback en lugar "
        f"de responder con contexto irrelevante. En cuanto a errores, la tasa de fallos de API real "
        f"es del {err['tasa_error_api_pct']}% ({err['errores_api']} de {err['total_ejecuciones']} "
        f"ejecuciones, propios del nivel gratuito de GitHub Models); adicionalmente se bloquearon "
        f"{err['bloqueados_seguridad']} consultas adversariales en la capa de seguridad.",
    )
    imagen(doc, "05_similitud_categoria.png",
           "Figura 4. Score de similitud (top-1) por categoría de consulta.", 14.0)
    imagen(doc, "04_tasa_errores.png",
           "Figura 5. Tasa de error global y distribución por tipo (API vs. seguridad).", 15.0)

    parrafo(doc, "4.4 Hallazgos clave (IE3, IE4)", size=10.5).runs[0].bold = True
    for hh in d["hallazgos"]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(hh)
        run.font.size = Pt(9.5)

    # ── 5. Seguridad ─────────────────────────────────────────────────────────
    h(doc, "5. Seguridad y uso responsable (IE6)")
    parrafo(
        doc,
        "Se integraron controles de seguridad y uso responsable en el diseño: "
        "(1) sanitización de inputs y detección de prompt injection mediante patrones, alineada "
        "al OWASP Top 10 for LLM Applications (LLM01); (2) gestión de la API key exclusivamente "
        "por variable de entorno (.env fuera del control de versiones); (3) anonimización de "
        "datos personales (RUT, correo, teléfono) antes de escribir los logs, en línea con la "
        "Ley N° 19.628 sobre protección de la vida privada en Chile; (4) truncado de la consulta "
        "registrada para minimizar datos; y (5) rate limiting por ventana deslizante para proteger "
        "la cuota de la API y mitigar abuso. Las consultas adversariales del dataset fueron "
        "bloqueadas antes de alcanzar el agente, sin consumir tokens. Desde una perspectiva ética "
        "y normativa, el agente opera sobre normativa portuaria sensible: la observabilidad debe "
        "equilibrar trazabilidad con privacidad, evitando registrar información operacional o "
        "personal identificable de los trabajadores del puerto.",
    )

    # ── 6. Recomendaciones ───────────────────────────────────────────────────
    h(doc, "6. Recomendaciones de optimización (IE7)")
    recomendaciones = [
        ("Reducir la latencia y el costo del LLM (cuello de botella). ",
         "Acotar max_iterations y la longitud del scratchpad ReAct, activar streaming para mejorar "
         "la latencia percibida, y cachear respuestas a consultas frecuentes. Evaluar una ruta "
         "directa (sin ReAct) para preguntas simples de una sola herramienta."),
        ("Umbral de relevancia con fallback. ",
         "Si el score top-1 < 0.45, responder explícitamente que la consulta está fuera del corpus "
         "o derivar a la búsqueda externa, evitando alucinaciones sobre contexto irrelevante."),
        ("Resiliencia ante errores de API. ",
         "Implementar reintentos con backoff exponencial y timeouts configurables para los "
         "errores 429/503 y los timeouts que inflan el p95; degradar de forma elegante."),
        ("Estabilizar la consistencia. ",
         "Bajar la temperature para consultas normativas (donde se busca determinismo) y fijar las "
         "respuestas a incidentes para reducir la variabilidad observada entre repeticiones."),
        ("Escalabilidad y sostenibilidad. ",
         "Exportar las métricas a un backend de observabilidad (p. ej. Grafana/Prometheus o "
         "LangSmith) y definir alertas sobre p95 de latencia, tasa de error y consumo de tokens "
         "para operación en producción."),
    ]
    for titulo_r, cuerpo_r in recomendaciones:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        rt = p.add_run(titulo_r)
        rt.bold = True
        rt.font.size = Pt(9.5)
        rc = p.add_run(cuerpo_r)
        rc.font.size = Pt(9.5)

    # ── 7. Conclusiones ──────────────────────────────────────────────────────
    h(doc, "7. Conclusiones")
    parrafo(
        doc,
        "La instrumentación permitió observar de punta a punta el comportamiento del agente y "
        "tomar decisiones basadas en datos. El análisis identificó al LLM como el cuello de "
        "botella dominante de latencia y costo, mostró que el score de similitud es una señal "
        "fiable para discriminar consultas fuera de dominio, y verificó que la capa de seguridad "
        "neutraliza intentos de prompt injection. Las recomendaciones propuestas apuntan a una "
        "solución más eficiente, robusta, segura y escalable, en línea con los objetivos de "
        "sostenibilidad del agente en producción.",
    )

    # ── 8. Referencias ───────────────────────────────────────────────────────
    h(doc, "8. Referencias (APA)")
    refs = [
        "Biblioteca del Congreso Nacional de Chile. (1999). Ley N° 19.628: Sobre protección de la vida privada. https://www.bcn.cl/leychile/navegar?idNorma=141599",
        "Chase, H. (2022). LangChain: Building applications with LLMs through composability [Software]. https://python.langchain.com",
        "Chroma. (2024). Chroma: The open-source embedding database. https://docs.trychroma.com",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., … Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459–9474.",
        "OpenAI. (2024). GPT-4o mini: Advancing cost-efficient intelligence. https://openai.com",
        "OWASP Foundation. (2023). OWASP Top 10 for Large Language Model Applications. https://owasp.org/www-project-top-10-for-large-language-model-applications/",
        "Sridharan, C. (2018). Distributed systems observability: A guide to building robust systems. O'Reilly Media.",
        "Streamlit Inc. (2024). Streamlit documentation. https://docs.streamlit.io",
        "Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K., & Cao, Y. (2023). ReAct: Synergizing reasoning and acting in language models. International Conference on Learning Representations (ICLR).",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        run = p.add_run(ref)
        run.font.size = Pt(8.5)

    doc.save(SALIDA)
    print(f"✅ Informe generado: {SALIDA}")


if __name__ == "__main__":
    construir(cargar_hallazgos())
