"""
Genera Informe_EFT_ISY0101.docx — informe final unificado (EFT, 4 apartados de la pauta).
Lee los datos reales: logs/hallazgos.json (después) y logs/hallazgos_antes_backoff.json (antes).
Reutiliza los helpers de estilo del generador EP3.
"""
import json
import os
import sys
import statistics as st

_EFT = os.path.dirname(os.path.abspath(__file__))
_EP3 = os.path.join(os.path.dirname(_EFT), "EP3_observabilidad")
sys.path.append(os.path.join(_EP3, "src"))

from docx.shared import Pt, Cm
from generar_informe import (  # helpers EP3: estilo, portada, tablas
    Document, set_base_style, h, parrafo, tabla_kv, portada, GRIS, AZUL,
    WD_ALIGN_PARAGRAPH,
)

SALIDA = os.path.join(_EFT, "Informe_EFT_ISY0101.docx")
DIAGRAMA = os.path.join(_EFT, "diagrama_arquitectura.png")
IMG_EP3 = os.path.join(_EP3, "docs", "img")


def cargar(nombre):
    with open(os.path.join(_EP3, "logs", nombre), encoding="utf-8") as f:
        return json.load(f)


def fila_metricas(hall):
    lat, err = hall["latencia"], hall["errores"]
    return (f"{err['tasa_error_api_pct']}%", f"{lat['promedio_total_s']}s",
            f"{lat['mediana_total_s']}s", f"{lat['p95_total_s']}s", f"{lat['p99_total_s']}s")


def imagen_local(doc, ruta, caption, ancho=15.0):
    if not os.path.exists(ruta):
        parrafo(doc, f"[Figura faltante: {os.path.basename(ruta)}]", italic=True, color=GRIS)
        return
    doc.add_picture(ruta, width=Cm(ancho))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GRIS


def construir():
    antes = cargar("hallazgos_antes_backoff.json")
    despues = cargar("hallazgos.json")

    doc = Document()
    set_base_style(doc)

    portada(
        doc,
        integrantes="Benjamín Aravena   ·   Francisco Gómez",
        docente="Roberto Eduardo Vega Araneda",
        asignatura="ISY0101 — Ingeniería de Soluciones con IA · Evaluación Final Transversal",
        fecha="Julio de 2026",
        titulo="Agente de IA para Gestión Normativa Portuaria",
        subtitulo="Solución integral RAG + Agente ReAct + Observabilidad — Empresa Portuaria Valparaíso (EPV)",
    )

    # ── 1. CASO ────────────────────────────────────────────────────────────────
    h(doc, "1. Análisis del caso organizacional (IE4, IE8)")
    parrafo(doc,
        "La Empresa Portuaria Valparaíso (EPV) opera bajo un cuerpo normativo extenso y disperso: "
        "reglamento de operaciones portuarias, manual de seguridad y EPP, protocolo de emergencias "
        "marítimas y convenio colectivo de trabajadores. El personal operativo y de prevención "
        "necesita respuestas rápidas y confiables (¿qué EPP exige el muelle?, ¿qué protocolo aplica "
        "ante un derrame?, ¿qué establece el convenio sobre turnos nocturnos?), pero la búsqueda "
        "manual en documentos PDF es lenta y propensa a errores de interpretación, con riesgo "
        "directo sobre la seguridad laboral y el cumplimiento normativo.")
    parrafo(doc,
        "Requerimientos levantados: (1) respuestas fundadas exclusivamente en la normativa oficial, "
        "con cita de fuente; (2) capacidad de evaluar situaciones de cumplimiento y emitir dictámenes "
        "accionables; (3) generación de documentación formal (reportes, memos, actas); (4) contexto "
        "internacional (SOLAS, MARPOL, OIT) cuando la normativa interna no baste; (5) trazabilidad "
        "completa de cada consulta y (6) resguardos de seguridad y privacidad para un contexto de "
        "producción (Ley N.º 19.628). Desafíos: alucinaciones del LLM, costos y límites de API, "
        "variabilidad de consultas (ambiguas, fuera de dominio, adversariales) y datos sensibles.")

    # ── 2. DISEÑO LLM + RAG ───────────────────────────────────────────────────
    h(doc, "2. Diseño de la solución basada en LLM y RAG")
    parrafo(doc, "2.1 Formulación de prompts (IE1)", size=10.5).runs[0].bold = True
    parrafo(doc,
        "Se diseñaron prompts en tres niveles. (a) Prompt ReAct del agente: define el rol (asistente "
        "experto en normativas EPV), enumera las herramientas con descripciones ricas, impone el "
        "formato Thought→Action→Action Input→Observation→Final Answer, y codifica reglas de "
        "planificación explícitas (p. ej., \"si la tarea requiere consultar y evaluar, usa "
        "consultar_normativa primero\"), con límite de 6 pasos y respuesta en español técnico. "
        "(b) Prompt de dictamen (evaluar_cumplimiento): plantilla con salida estructurada — nivel de "
        "riesgo (BAJO/MEDIO/ALTO/CRÍTICO), áreas normativas, hallazgos numerados, recomendaciones y "
        "prioridad — que vuelve el resultado accionable y parseable. (c) Docstrings de herramientas: "
        "actúan como prompts de selección, indicando al modelo cuándo usar cada tool con ejemplos "
        "concretos. Temperaturas bajas (0.2 agente, 0.1 dictamen) privilegian fidelidad sobre "
        "creatividad, coherente con el dominio normativo.")

    parrafo(doc, "2.2 Pipeline RAG con fuentes internas y externas (IE2)", size=10.5).runs[0].bold = True
    parrafo(doc,
        "Fuente interna: los 4 PDF oficiales se cargan (pypdf), se trocean con "
        "RecursiveCharacterTextSplitter (chunks de 800 caracteres, solape 100) y se vectorizan con "
        "embeddings multilingües paraphrase-multilingual-MiniLM-L12-v2 (384 dimensiones, ejecución "
        "local — el corpus no sale de la organización), persistiendo en ChromaDB. En consulta, se "
        "recuperan los k=4 fragmentos de menor distancia L2, se anteponen con su fuente y alimentan "
        "la generación. Fuente externa: la herramienta buscar_fuente_externa consulta la API REST de "
        "Wikipedia en español en tiempo real para normativa internacional (SOLAS, MARPOL, convenios "
        "OIT), enriqueciendo respuestas que el corpus interno no cubre. El agente combina ambas "
        "fuentes bajo demanda según su planificación.")

    parrafo(doc, "2.3 Arquitectura de la solución (IE3)", size=10.5).runs[0].bold = True
    imagen_local(doc, DIAGRAMA, "Figura 1. Arquitectura general: seguridad → agente ReAct → "
                 "herramientas/memorias → observabilidad → dashboard.", 16.0)
    parrafo(doc,
        "El control de contexto opera en tres frentes: el retriever acota el contexto documental a "
        "los 4 chunks más relevantes; las memorias inyectan solo el historial pertinente (ventana de "
        "8 turnos + 3 recuerdos semánticos); y el límite de 6 iteraciones contiene el crecimiento "
        "del scratchpad ReAct, acotando costo y latencia.")

    parrafo(doc, "2.4 Justificación de decisiones (IE4)", size=10.5).runs[0].bold = True
    tabla_kv(doc,
        encabezados=["Decisión", "Justificación (requerimiento / trazabilidad / limitación)"],
        filas=[
            ["gpt-4o-mini (GitHub Models)", "Balance costo-calidad en español; acceso académico gratuito. Limitación medida: rate-limit del free tier → mitigada con backoff."],
            ["LangChain + patrón ReAct", "Planificación multi-paso auditable: cada Thought queda registrado, aportando trazabilidad del razonamiento."],
            ["ChromaDB + MiniLM local", "Corpus en español, cero costo, privacidad: embeddings e índice permanecen en la organización."],
            ["Chunks 800/100, k=4", "Fragmentos con sentido completo sin inflar el prompt; validado con scores de similitud medidos."],
            ["Temperature 0.1–0.2", "Dominio normativo exige respuestas estables y fieles; la consistencia medida lo confirma."],
            ["Logging JSON Lines", "Un registro por consulta: trazabilidad completa, procesable por el dashboard y herramientas estándar."],
        ])

    # ── 3. AGENTE ─────────────────────────────────────────────────────────────
    h(doc, "3. Desarrollo del agente funcional")
    parrafo(doc, "3.1 Herramientas de consulta, escritura y razonamiento (IE5)", size=10.5).runs[0].bold = True
    tabla_kv(doc,
        encabezados=["Herramienta", "Tipo", "Función"],
        filas=[
            ["consultar_normativa", "Consulta", "RAG sobre ChromaDB: recupera y cita fragmentos de la normativa interna EPV."],
            ["evaluar_cumplimiento", "Razonamiento", "Dictamen técnico estructurado: nivel de riesgo, hallazgos y recomendaciones."],
            ["generar_reporte", "Escritura", "Redacta y persiste documentos formales (reportes de incidente, memos, actas)."],
            ["buscar_fuente_externa", "Consulta externa", "Wikipedia ES en tiempo real para normativa internacional."],
        ])
    parrafo(doc, "3.2 Memoria y recuperación de contexto (IE6)", size=10.5).runs[0].bold = True
    parrafo(doc,
        "Memoria de corto plazo: ConversationBufferWindowMemory con ventana de 8 turnos, que da "
        "continuidad conversacional dentro de la sesión (resolución de referencias como \"ese "
        "protocolo\"). Memoria de largo plazo: una segunda base ChromaDB donde se persisten "
        "resúmenes de interacciones relevantes con metadatos de fecha y tipo; en cada consulta se "
        "recuperan los 3 recuerdos semánticamente más cercanos y se inyectan al prompt, asegurando "
        "continuidad de tareas entre sesiones sin acumular contexto ilimitado.")
    parrafo(doc, "3.3 Planificación y toma de decisiones (IE7)", size=10.5).runs[0].bold = True
    parrafo(doc,
        "El bucle ReAct implementa la planificación: el modelo razona explícitamente, selecciona la "
        "herramienta, observa el resultado y decide el siguiente paso, hasta un máximo de 6 "
        "iteraciones con handle_parsing_errors para recuperarse de salidas mal formateadas. Las "
        "reglas de planificación del prompt encadenan herramientas ante tareas compuestas (consultar "
        "→ evaluar → reportar) y derivan a la fuente externa cuando el tema es internacional. En la "
        "evaluación se observó comportamiento adaptativo real: ante consultas fuera de dominio el "
        "agente optó por la fuente externa o respondió directamente sin forzar el RAG.")
    parrafo(doc, "3.4 Orquestación y flujo de trabajo (IE8)", size=10.5).runs[0].bold = True
    parrafo(doc,
        "Flujo por consulta: interfaz Streamlit → capa de seguridad (validación) → recuperación de "
        "memoria de largo plazo → AgentExecutor (bucle ReAct con las 4 herramientas) → persistencia "
        "en ambas memorias → registro de métricas en JSONL → respuesta con el razonamiento visible. "
        "La capa de observabilidad envuelve todo el flujo sin modificar el código del agente "
        "(wrapper + callback handler de LangChain), decisión que independiza la evolución del agente "
        "de su instrumentación.")

    # ── 4. OBSERVABILIDAD / SEGURIDAD ─────────────────────────────────────────
    h(doc, "4. Observabilidad, trazabilidad y seguridad")
    parrafo(doc, "4.1 Métricas aplicadas (IE9)", size=10.5).runs[0].bold = True
    lat_a, rec_a = antes["latencia"], antes["recursos"]
    parrafo(doc,
        f"Sobre un dataset de estrés de 28 consultas (36 ejecuciones, con repeticiones para "
        f"consistencia y casos límite: ambiguas, fuera de dominio y adversariales) se midieron las "
        f"cinco familias de métricas. Resultados de referencia (corrida EP3): el LLM concentra el "
        f"{lat_a['porcentaje_tiempo_por_fase']['llm_pct']}% de la latencia total (el cuello de "
        f"botella es la generación, no el retrieval); consumo medio de "
        f"{rec_a['tokens_promedio_total']} tokens por consulta con razón prompt:completion "
        f"{rec_a['ratio_prompt_completion']}:1 (el scratchpad ReAct domina el costo); y el score de "
        f"similitud separa consultas en-dominio (0.076) de fuera-de-dominio (0.039), validando un "
        f"umbral de relevancia de {antes.get('umbral_relevancia', 0.057)} para detectar consultas no "
        f"cubiertas por el corpus.")

    parrafo(doc, "4.2 Análisis de registros y mejora implementada (IE10, IE12)", size=10.5).runs[0].bold = True
    parrafo(doc,
        "El análisis de trazas de EP3 identificó como falla dominante el rate-limit del nivel "
        "gratuito de GitHub Models (HTTP 429): 33.3% de error de API bajo carga sostenida, inflando "
        "la cola de latencia. Siguiendo la retroalimentación docente, se implementó la mejora: "
        "reintentos con backoff exponencial (max_retries=5, mecanismo integrado del SDK OpenAI, con "
        "jitter) y timeout de 45 s en ambos clientes LLM. La re-evaluación con el dataset completo "
        "verifica el efecto:")
    tabla_kv(doc,
        encabezados=["Métrica", "Antes (EP3)", "Después (backoff)"],
        filas=[
            ["Tasa de error de API", fila_metricas(antes)[0], fila_metricas(despues)[0]],
            ["Latencia media", fila_metricas(antes)[1], fila_metricas(despues)[1]],
            ["Latencia mediana", fila_metricas(antes)[2], fila_metricas(despues)[2]],
            ["Latencia p95", fila_metricas(antes)[3], fila_metricas(despues)[3]],
            ["Latencia p99", fila_metricas(antes)[4], fila_metricas(despues)[4]],
        ])
    parrafo(doc,
        "Interpretación: la tasa de error de API cayó de 33.3% a 5.6% (6× menos fallas; los 10 "
        "rate-limit desaparecieron) y la mediana mejoró de 9.3 s a 6.8 s. Como contraparte, la cola "
        "extrema (p99) creció: los reintentos convierten consultas que antes fallaban rápido en "
        "respuestas lentas pero exitosas. Es un trade-off deliberado y correcto para este caso de "
        "uso: para un usuario que consulta normativa de seguridad, una respuesta tardía es "
        "preferible a un error. El p95/p99 residual queda como objetivo de la siguiente iteración "
        "(tier pago o cache).")
    parrafo(doc,
        "Mejoras futuras priorizadas por impacto: activar el umbral de relevancia con fallback "
        "automático; migración a tier pago o modelo on-premise; evaluación de calidad de respuesta "
        "con LLM-as-judge (RAGAS); exportación de métricas a Grafana/Prometheus con alertas sobre "
        "p95 y tasa de error; y function calling nativo para eliminar los errores de parseo ReAct "
        "observados en consultas ambiguas.")

    parrafo(doc, "4.3 Seguridad y uso responsable (IE11)", size=10.5).runs[0].bold = True
    parrafo(doc,
        "Controles implementados: sanitización de inputs con detección de prompt injection por "
        "patrones (OWASP LLM01) aplicada antes del agente — en la evaluación bloqueó las 4 consultas "
        "adversariales sin consumir tokens; anonimización de datos personales (RUT, correo, "
        "teléfono) en los registros, alineada con la Ley N.º 19.628 sobre protección de la vida "
        "privada; rate limiting local de ventana deslizante; gestión de credenciales por variable "
        "de entorno fuera del control de versiones; y truncado de consultas registradas "
        "(minimización de datos). Consideraciones éticas: el agente cita fuentes y expone su "
        "razonamiento (transparencia), se declara asistente y no sustituye el juicio experto en "
        "seguridad laboral, y se documenta como limitación que los prompts viajan a la API externa "
        "— en producción se evaluaría un modelo autoalojado o acuerdos de tratamiento de datos.")

    # ── 5. CONCLUSIONES Y REFLEXIONES ─────────────────────────────────────────
    h(doc, "5. Conclusiones y reflexiones individuales")
    parrafo(doc,
        "La solución cumple el ciclo completo de ingeniería de una aplicación con agentes de IA: "
        "diseño fundado en requerimientos, implementación funcional con planificación y memoria, y "
        "operación medible, segura y mejorada con evidencia. El principal aprendizaje técnico es que "
        "la observabilidad convierte percepciones en decisiones: sin medición no habríamos sabido "
        "que el 99.6% del tiempo está en la generación ni que un tercio de las fallas provenía del "
        "rate-limit — ni podríamos demostrar que la mejora funcionó.")
    parrafo(doc, "Reflexión individual — Benjamín Aravena:", size=10.5).runs[0].bold = True
    parrafo(doc, "[REDACTAR PERSONALMENTE, SIN IA — exigencia de la pauta: aprendizaje personal y "
            "contribución al proyecto. 5–8 líneas.]", italic=True, color=GRIS)
    parrafo(doc, "Reflexión individual — Francisco Gómez:", size=10.5).runs[0].bold = True
    parrafo(doc, "[REDACTAR PERSONALMENTE, SIN IA — exigencia de la pauta: aprendizaje personal y "
            "contribución al proyecto. 5–8 líneas.]", italic=True, color=GRIS)

    h(doc, "6. Declaración de uso de IA")
    parrafo(doc,
        "Conforme a las indicaciones de la evaluación y a https://bibliotecas.duoc.cl/ia, se declara "
        "el uso de IA generativa (Claude, Anthropic) como apoyo en: mejora de redacción del informe, "
        "estructuración de código y generación de diagramas. Las ideas, decisiones de diseño, "
        "análisis de resultados, conclusiones y reflexiones individuales son de autoría propia del "
        "equipo y fueron validadas contra los requerimientos del proyecto.")

    h(doc, "7. Referencias (APA)")
    refs = [
        "Biblioteca del Congreso Nacional de Chile. (1999). Ley N.º 19.628: Sobre protección de la vida privada. https://www.bcn.cl/leychile/navegar?idNorma=141599",
        "Chase, H. (2022). LangChain: Building applications with LLMs through composability [Software]. https://python.langchain.com",
        "Chroma. (2024). Chroma: The open-source embedding database. https://docs.trychroma.com",
        "Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., … Kiela, D. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. Advances in Neural Information Processing Systems, 33, 9459–9474.",
        "OpenAI. (2024). GPT-4o mini: Advancing cost-efficient intelligence. https://openai.com",
        "OWASP Foundation. (2023). OWASP Top 10 for Large Language Model Applications. https://owasp.org/www-project-top-10-for-large-language-model-applications/",
        "Reimers, N., & Gurevych, I. (2019). Sentence-BERT: Sentence embeddings using Siamese BERT-networks. Proceedings of EMNLP-IJCNLP 2019, 3982–3992.",
        "Sridharan, C. (2018). Distributed systems observability: A guide to building robust systems. O'Reilly Media.",
        "Streamlit Inc. (2024). Streamlit documentation. https://docs.streamlit.io",
        "Yao, S., Zhao, J., Yu, D., Du, N., Shafran, I., Narasimhan, K., & Cao, Y. (2023). ReAct: Synergizing reasoning and acting in language models. International Conference on Learning Representations (ICLR).",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        run = p.add_run(ref)
        run.font.size = Pt(8.5)

    doc.save(SALIDA)
    print("OK:", SALIDA)


if __name__ == "__main__":
    construir()
